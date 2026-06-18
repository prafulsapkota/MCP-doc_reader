import os
import math
from typing import Optional, Generator
import docx
from fastmcp import FastMCP
from src.utils import parse_page_range, stream_text_by_sentences

def register_docx_tools(mcp: FastMCP):
    @mcp.tool(name="read_docx", description="Read Word DOCX files and convert to Markdown. Optional parameter page_range is supported as segment range (every ~30 paragraphs as a logical page).")
    def read_docx(file_path: str, page_range: Optional[str] = None) -> str:
        """Reads DOCX and outputs Markdown."""
        if not os.path.exists(file_path):
            return f"**Error:** File not found at `{file_path}`"
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            page_size = 30
            total_pages = math.ceil(len(paragraphs) / page_size) if paragraphs else 1
            
            pages_to_read = parse_page_range(page_range, total_pages)
            
            output = []
            output.append(f"# Word Document: {os.path.basename(file_path)}")
            output.append(f"*Estimated Pages (30 paras/page):* {total_pages}\n---")
            
            for p_idx in pages_to_read:
                start_idx = p_idx * page_size
                end_idx = min(start_idx + page_size, len(paragraphs))
                page_paras = paragraphs[start_idx:end_idx]
                
                output.append(f"## Section/Page {p_idx + 1}")
                for para in page_paras:
                    output.append(para + "\n")
                    
            return "\n".join(output)
        except Exception as e:
            return f"**Error reading DOCX:** {str(e)}"

    @mcp.tool(name="stream_docx", description="Stream DOCX text sentence-by-sentence.")
    def stream_docx(file_path: str, page_range: Optional[str] = None) -> Generator[str, None, None]:
        """Streams DOCX content sentence-by-sentence."""
        if not os.path.exists(file_path):
            yield f"**Error:** File not found at `{file_path}`"
            return
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            page_size = 30
            total_pages = math.ceil(len(paragraphs) / page_size) if paragraphs else 1
            
            pages_to_read = parse_page_range(page_range, total_pages)
            
            yield f"# Streaming DOCX: {os.path.basename(file_path)}\n"
            
            for p_idx in pages_to_read:
                yield f"\n## Section/Page {p_idx + 1}\n"
                start_idx = p_idx * page_size
                end_idx = min(start_idx + page_size, len(paragraphs))
                page_paras = paragraphs[start_idx:end_idx]
                
                for para in page_paras:
                    for sentence in stream_text_by_sentences(para):
                         yield sentence
        except Exception as e:
            yield f"**Error streaming DOCX:** {str(e)}"

    @mcp.tool(name="read_docx_data", description="Read Word DOCX file passed as base64 encoded string. Optional file_name is for context. Optional page_range is supported as segment range (every ~30 paragraphs as a logical page).")
    def read_docx_data(file_content_b64: str, file_name: Optional[str] = None, page_range: Optional[str] = None) -> str:
        """Reads base64-encoded DOCX and outputs Markdown."""
        try:
            from src.utils import decode_b64_to_stream
            stream = decode_b64_to_stream(file_content_b64)
            doc = docx.Document(stream)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            page_size = 30
            total_pages = math.ceil(len(paragraphs) / page_size) if paragraphs else 1
            
            pages_to_read = parse_page_range(page_range, total_pages)
            
            display_name = file_name if file_name else "Uploaded Document"
            output = []
            output.append(f"# Word Document: {display_name}")
            output.append(f"*Estimated Pages (30 paras/page):* {total_pages}\n---")
            
            for p_idx in pages_to_read:
                start_idx = p_idx * page_size
                end_idx = min(start_idx + page_size, len(paragraphs))
                page_paras = paragraphs[start_idx:end_idx]
                
                output.append(f"## Section/Page {p_idx + 1}")
                for para in page_paras:
                    output.append(para + "\n")
                    
            return "\n".join(output)
        except Exception as e:
            return f"**Error reading DOCX data:** {str(e)}"

    @mcp.tool(name="stream_docx_data", description="Stream DOCX text sentence-by-sentence from base64 encoded string.")
    def stream_docx_data(file_content_b64: str, file_name: Optional[str] = None, page_range: Optional[str] = None) -> Generator[str, None, None]:
        """Streams DOCX content sentence-by-sentence from base64 encoded string."""
        try:
            from src.utils import decode_b64_to_stream
            stream = decode_b64_to_stream(file_content_b64)
            doc = docx.Document(stream)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            page_size = 30
            total_pages = math.ceil(len(paragraphs) / page_size) if paragraphs else 1
            
            pages_to_read = parse_page_range(page_range, total_pages)
            
            display_name = file_name if file_name else "Uploaded Document"
            yield f"# Streaming DOCX: {display_name}\n"
            
            for p_idx in pages_to_read:
                yield f"\n## Section/Page {p_idx + 1}\n"
                start_idx = p_idx * page_size
                end_idx = min(start_idx + page_size, len(paragraphs))
                page_paras = paragraphs[start_idx:end_idx]
                
                for para in page_paras:
                    for sentence in stream_text_by_sentences(para):
                        yield sentence
        except Exception as e:
            yield f"**Error streaming DOCX data:** {str(e)}"
